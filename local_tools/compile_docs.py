#ident "@(#)$Format:LocalFoodAI:app.py:%an:%ae:%ad:%cn:%ce:%cd:%H:%D:%N$"
import os
import sys
import re
from datetime import datetime

# Ensure stdout handles UTF-8 correctly
if sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout.reconfigure(encoding='utf-8')

# ReportLab imports
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak, Table, TableStyle, Preformatted
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors

# PyMuPDF import
import fitz

# Keywords to track for the Page-Search Index
INDEX_KEYWORDS = ["MySQL", "Docker", "Ollama", "Telemetry", "RAG", "SearXNG", "Zabbix", "Airflow", "Kubernetes", "Jenkins", "Longhorn", "MetalLB", "WSL"]

def clean_text(text):
    """Strips characters that cannot be encoded in standard PDF fonts (Latin-1)."""
    cleaned = []
    for char in text:
        try:
            # Built-in PDF Helvetica/Courier supports Latin-1 character set
            char.encode('latin-1')
            cleaned.append(char)
        except UnicodeEncodeError:
            # Skip non-latin-1 characters (like emojis)
            continue
    return "".join(cleaned)

class PDFCompiler:
    def __init__(self, md_path, pdf_path):
        self.md_path = md_path
        self.pdf_path = pdf_path
        self.doc_title = os.path.splitext(os.path.basename(md_path))[0].replace("_", " ").title()
        self.author = "Roni Fernandes Dias / François LANGE"
        self.date_str = datetime.now().strftime("%B %d, %Y")
        self.toc_map = {}
        self.index_map = {kw: set() for kw in INDEX_KEYWORDS}
        self.links = [] # List of tuples: (Link Text, URL)

    def parse_markdown(self):
        """Parses markdown file and returns raw lines and structural elements."""
        with open(self.md_path, 'r', encoding='utf-8', errors='ignore') as f:
            content = f.read()

        # Extract markdown links [Text](URL)
        link_pattern = r'\[([^\]]+)\]\((https?://[^\)]+)\)'
        self.links = re.findall(link_pattern, content)

        # Split into lines
        lines = content.split('\n')
        # Filter out #ident lines
        filtered_lines = [l for l in lines if not l.startswith('#ident') and not '#ident' in l]
        return filtered_lines

    def build_flowables(self, lines, pass_num=1):
        """Converts raw markdown lines into ReportLab flowables."""
        styles = getSampleStyleSheet()
        
        # Modify existing styles to use standard Helvetica fonts
        styles['Normal'].fontName = 'Helvetica'
        styles['Normal'].fontSize = 10
        styles['Normal'].leading = 14
        styles['Normal'].textColor = colors.HexColor("#2C3E50")
        
        styles['Heading1'].fontName = 'Helvetica-Bold'
        styles['Heading1'].fontSize = 18
        styles['Heading1'].leading = 22
        styles['Heading1'].textColor = colors.HexColor("#1A365D")
        styles['Heading1'].spaceAfter = 12
        styles['Heading1'].keepWithNext = True

        styles['Heading2'].fontName = 'Helvetica-Bold'
        styles['Heading2'].fontSize = 13
        styles['Heading2'].leading = 17
        styles['Heading2'].textColor = colors.HexColor("#2B6CB0")
        styles['Heading2'].spaceAfter = 8
        styles['Heading2'].keepWithNext = True

        # Custom Code block style
        code_style = ParagraphStyle(
            'CodeBlock',
            fontName='Courier',
            fontSize=8,
            leading=10,
            textColor=colors.HexColor("#1A202C"),
            backColor=colors.HexColor("#EDF2F7"),
            borderPadding=8,
            borderWidth=0.5,
            borderColor=colors.HexColor("#CBD5E0"),
            spaceAfter=10
        )

        flowables = []

        # --- 1. COVER PAGE ---
        flowables.append(Spacer(1, 150))
        title_style = ParagraphStyle(
            'CoverTitle',
            fontName='Helvetica-Bold',
            fontSize=26,
            leading=32,
            alignment=1, # Centered
            textColor=colors.HexColor("#1A365D"),
            spaceAfter=30
        )
        meta_style = ParagraphStyle(
            'CoverMeta',
            fontName='Helvetica',
            fontSize=11,
            leading=16,
            alignment=1, # Centered
            textColor=colors.HexColor("#718096")
        )
        flowables.append(Paragraph(clean_text(self.doc_title), title_style))
        flowables.append(Spacer(1, 100))
        flowables.append(Paragraph(f"<b>Project Category:</b> B1AI DEVOP1<br/>"
                                   f"<b>Author:</b> {self.author}<br/>"
                                   f"<b>Date:</b> {self.date_str}", meta_style))
        flowables.append(PageBreak())

        # --- 2. BLANK PAGE ---
        flowables.append(Spacer(1, 1))
        flowables.append(PageBreak())

        # --- 3. TABLE OF CONTENTS (PASS 2 ONLY) ---
        if pass_num == 2 and self.toc_map:
            toc_title_style = ParagraphStyle(
                'TOCTitle',
                parent=styles['Heading1'],
                keepWithNext=True
            )
            flowables.append(Paragraph("Table of Contents", toc_title_style))
            flowables.append(Spacer(1, 15))
            
            toc_item_style = ParagraphStyle(
                'TOCItem',
                fontName='Helvetica',
                fontSize=10,
                leading=14,
                textColor=colors.HexColor("#2D3748")
            )
            
            for heading, page in self.toc_map.items():
                dots_count = max(4, 75 - len(heading))
                dots = '.' * dots_count
                toc_line = f"<b>{clean_text(heading)}</b> {dots} {page}"
                flowables.append(Paragraph(toc_line, toc_item_style))
                flowables.append(Spacer(1, 6))
                
            flowables.append(PageBreak())

        # --- 4. BODY CONTENT ---
        in_code_block = False
        code_lines = []

        for line in lines:
            stripped = line.strip()
            
            # Code block toggles
            if stripped.startswith("```"):
                if in_code_block:
                    # End code block - Wrap long lines to prevent PDF overflow
                    wrapped_code_lines = []
                    for cl in code_lines:
                        if len(cl) > 75:
                            indent_len = len(cl) - len(cl.lstrip())
                            indent_str = cl[:indent_len]
                            content = cl[indent_len:]
                            limit = 75 - indent_len
                            if limit <= 10:
                                limit = 75
                                indent_str = ""
                                content = cl
                            import textwrap
                            wrapped = textwrap.wrap(content, width=limit, break_long_words=True, break_on_hyphens=False)
                            if wrapped:
                                wrapped_code_lines.append(indent_str + wrapped[0])
                                for w in wrapped[1:]:
                                    wrapped_code_lines.append(indent_str + "    " + w)
                            else:
                                wrapped_code_lines.append(cl)
                        else:
                            wrapped_code_lines.append(cl)
                    code_text = "\n".join(wrapped_code_lines)
                    flowables.append(Preformatted(clean_text(code_text), code_style))
                    code_lines = []
                    in_code_block = False
                else:
                    # Start code block
                    in_code_block = True
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            # Headings
            if stripped.startswith("# "):
                heading_text = stripped[2:]
                flowables.append(Paragraph(clean_text(heading_text), styles['Heading1']))
            elif stripped.startswith("## "):
                heading_text = stripped[3:]
                flowables.append(Paragraph(clean_text(heading_text), styles['Heading2']))
            elif stripped.startswith("### "):
                heading_text = stripped[4:]
                flowables.append(Paragraph(clean_text(heading_text), styles['Heading2']))
            
            # Bullet/Numbered list items
            elif stripped.startswith("- ") or stripped.startswith("* "):
                bullet_style = ParagraphStyle(
                    'BulletList',
                    parent=styles['Normal'],
                    leftIndent=15,
                    firstLineIndent=-10,
                    spaceAfter=4
                )
                flowables.append(Paragraph(f"• {clean_text(stripped[2:])}", bullet_style))
            
            # Table Row
            elif stripped.startswith("|") and stripped.endswith("|"):
                if "---" in stripped:
                    continue # Skip separator line
                cells = [c.strip() for c in stripped.split("|")[1:-1]]
                wrapped_cells = [Paragraph(clean_text(c), styles['Normal']) for c in cells]
                
                table_style = TableStyle([
                    ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#1A365D")),
                    ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                    ('ALIGN', (0,0), (-1,-1), 'LEFT'),
                    ('BOTTOMPADDING', (0,0), (-1,0), 6),
                    ('BACKGROUND', (0,1), (-1,-1), colors.HexColor("#F7FAFC")),
                    ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'),
                    ('TOPPADDING', (0,0), (-1,-1), 6),
                    ('BOTTOMPADDING', (0,0), (-1,-1), 6),
                ])
                
                col_width = 504.0 / max(1, len(cells))
                colWidths = [col_width] * len(cells)
                t = Table([wrapped_cells], colWidths=colWidths)
                t.setStyle(table_style)
                flowables.append(t)
                flowables.append(Spacer(1, 6))

            # Blank spacer
            elif not stripped:
                flowables.append(Spacer(1, 8))
            
            # Normal paragraph
            else:
                flowables.append(Paragraph(clean_text(stripped), styles['Normal']))

        # --- 5. AUTOMATED BIBLIOGRAPHY & REFERENCES (PASS 2 ONLY) ---
        if pass_num == 2 and self.links:
            flowables.append(Spacer(1, 20))
            flowables.append(Paragraph("References and Bibliography", styles['Heading1']))
            flowables.append(Spacer(1, 10))
            
            ref_data = [["Resource Description", "URL Endpoint"]]
            for text, url in self.links:
                ref_data.append([clean_text(text), clean_text(url)])
                
            ref_table = Table(ref_data, colWidths=[200, 300])
            ref_table.setStyle(TableStyle([
                ('BACKGROUND', (0,0), (-1,0), colors.HexColor("#2B6CB0")),
                ('TEXTCOLOR', (0,0), (-1,0), colors.white),
                ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor("#E2E8F0")),
                ('PADDING', (0,0), (-1,-1), 6),
                ('VALIGN', (0,0), (-1,-1), 'TOP'),
            ]))
            flowables.append(ref_table)
            flowables.append(PageBreak())

        # --- 6. PAGE-SEARCH INDEX (PASS 2 ONLY) ---
        if pass_num == 2 and any(self.index_map.values()):
            flowables.append(Paragraph("Page Search Index", styles['Heading1']))
            flowables.append(Spacer(1, 15))
            
            index_style = ParagraphStyle(
                'IndexItem',
                fontName='Helvetica',
                fontSize=10,
                leading=14,
                textColor=colors.HexColor("#2D3748")
            )
            
            for term in sorted(self.index_map.keys()):
                pages = sorted(list(self.index_map[term]))
                if pages:
                    pages_str = ", ".join(map(str, pages))
                    dots_count = max(4, 60 - len(term))
                    dots = '.' * dots_count
                    flowables.append(Paragraph(f"<b>{clean_text(term)}</b> {dots} Page(s): {pages_str}", index_style))
                    flowables.append(Spacer(1, 6))

        return flowables

    def draw_page_number_and_header(self, canvas, doc):
        """Callback to draw professional header and footer on content pages."""
        if doc.page <= 2:
            return
            
        canvas.saveState()
        
        # Header
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(colors.HexColor("#718096"))
        canvas.drawString(54, 750, f"Project DEVOP1  |  {clean_text(self.doc_title)}")
        canvas.setStrokeColor(colors.HexColor("#E2E8F0"))
        canvas.setLineWidth(0.5)
        canvas.line(54, 742, 558, 742)
        
        # Footer
        canvas.drawString(54, 40, f"Confidential  |  {self.author}")
        canvas.drawRightString(558, 40, f"Page {doc.page}")
        canvas.restoreState()

    def compile(self):
        """Runs the two-pass compilation pipeline."""
        print(f"Compiling document: {self.doc_title}...")
        
        lines = self.parse_markdown()

        # PASS 1: Build Draft PDF
        draft_pdf_path = self.pdf_path + ".draft.pdf"
        doc = SimpleDocTemplate(
            draft_pdf_path,
            pagesize=letter,
            leftMargin=54, rightMargin=54,
            topMargin=72, bottomMargin=72
        )
        
        flowables_draft = self.build_flowables(lines, pass_num=1)
        doc.build(flowables_draft, onFirstPage=self.draw_page_number_and_header, onLaterPages=self.draw_page_number_and_header)
        
        # Parse Draft PDF with PyMuPDF
        self.parse_draft_pdf(draft_pdf_path)
        
        if os.path.exists(draft_pdf_path):
            os.remove(draft_pdf_path)

        # PASS 2: Build Final PDF
        doc_final = SimpleDocTemplate(
            self.pdf_path,
            pagesize=letter,
            leftMargin=54, rightMargin=54,
            topMargin=72, bottomMargin=72
        )
        flowables_final = self.build_flowables(lines, pass_num=2)
        doc_final.build(flowables_final, onFirstPage=self.draw_page_number_and_header, onLaterPages=self.draw_page_number_and_header)
        print(f"[SUCCESS] Created: {self.pdf_path}")

    def parse_draft_pdf(self, draft_path):
        """Reads draft PDF and indexes text for TOC mapping and Keyword page searching."""
        doc = fitz.open(draft_path)
        
        lines = self.parse_markdown()
        for line in lines:
            stripped = line.strip()
            if stripped.startswith("# ") or stripped.startswith("## ") or stripped.startswith("### "):
                title = stripped.split(" ", 1)[1]
                for page_num in range(len(doc)):
                    page = doc[page_num]
                    instances = page.search_for(title)
                    if instances:
                        self.toc_map[title] = page_num + 1
                        break

        for keyword in INDEX_KEYWORDS:
            for page_num in range(len(doc)):
                page = doc[page_num]
                if page_num + 1 <= 2:
                    continue
                instances = page.search_for(keyword)
                if instances:
                    self.index_map[keyword].add(page_num + 1)
        doc.close()

def compile_to_docx(md_path, docx_path):
    import docx
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    
    doc = docx.Document()
    
    # Page setup - 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    doc_title = os.path.splitext(os.path.basename(md_path))[0].replace("_", " ").title()
    
    # Title heading
    title_p = doc.add_paragraph()
    title_run = title_p.add_run(doc_title)
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(24)
    title_run.font.bold = True
    title_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D) # Dark navy
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_after = Pt(24)
    
    with open(md_path, 'r', encoding='utf-8', errors='ignore') as f:
        content = f.read()
        
    lines = content.split('\n')
    filtered_lines = [l for l in lines if not l.startswith('#ident') and not '#ident' in l]
    
    in_code = False
    code_lines = []
    
    for line in filtered_lines:
        stripped = line.strip()
        
        # Code block toggle
        if stripped.startswith("```"):
            if in_code:
                # End code block - add as formatted pre text
                code_text = "\n".join(code_lines)
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.5)
                p.paragraph_format.space_after = Pt(10)
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(8.5)
                run.font.color.rgb = RGBColor(0x2D, 0x37, 0x48)
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
            
        if in_code:
            code_lines.append(line)
            continue
            
        # Headings
        if stripped.startswith("# "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(18)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)
            p.paragraph_format.space_before = Pt(18)
            p.paragraph_format.space_after = Pt(8)
            
        elif stripped.startswith("## "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[3:])
            run.font.name = 'Arial'
            run.font.size = Pt(14)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            p.paragraph_format.space_before = Pt(14)
            p.paragraph_format.space_after = Pt(6)
            
        elif stripped.startswith("### "):
            p = doc.add_paragraph()
            run = p.add_run(stripped[4:])
            run.font.name = 'Arial'
            run.font.size = Pt(12)
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)
            p.paragraph_format.space_before = Pt(12)
            p.paragraph_format.space_after = Pt(4)
            
        # Lists
        elif stripped.startswith("- ") or stripped.startswith("* "):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(stripped[2:])
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(4)
            
        # Tables
        elif stripped.startswith("|") and stripped.endswith("|"):
            if "---" in stripped:
                continue
            cells = [c.strip() for c in stripped.split("|")[1:-1]]
            
            table = None
            if len(doc.element.body) > 0 and doc.element.body[-1].tag.endswith('tbl'):
                table = doc.tables[-1]
                row = table.add_row()
            else:
                table = doc.add_table(rows=1, cols=len(cells))
                table.style = 'Table Grid'
                row = table.rows[0]
                
            for idx, cell_val in enumerate(cells):
                cell = row.cells[idx]
                cell.text = cell_val
                for p in cell.paragraphs:
                    p.paragraph_format.space_after = Pt(2)
                    p.paragraph_format.space_before = Pt(2)
                    for r in p.runs:
                        r.font.name = 'Arial'
                        r.font.size = Pt(9.5)
                        
        elif not stripped:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(6)
            
        else:
            p = doc.add_paragraph()
            run = p.add_run(stripped)
            run.font.name = 'Arial'
            run.font.size = Pt(10)
            p.paragraph_format.space_after = Pt(6)
            
    doc.save(docx_path)
    print(f"[SUCCESS] Created DOCX: {docx_path}")

def compile_all_documents():
    workspace = os.path.abspath(os.path.join(os.path.dirname(__file__), ".."))
    docs_dir = os.path.join(workspace, "docs")
    
    target_mappings = {
        # Root level
        "description.md": "description.pdf",
        "domainproject.md": "domainproject.pdf",
        "Project.md": "Project.pdf",
        "Retro_Planning.md": "Retro_Planning.pdf",
        "persona-template.md": "persona-template.pdf",
        "User stories.md": "User stories.pdf",
        "Wireframes.md": "Wireframes.pdf",
        "Presentation_Guide.md": "Presentation_Guide.pdf",
        
        # Docs level
        "docs/architecture.md": "docs/architecture.pdf",
        "docs/Backup_Procedure.md": "docs/Backup_Procedure.pdf",
        "docs/Data_Ingestion.md": "docs/Data_Ingestion.pdf",
        "docs/disaster_recovery_plan.md": "docs/disaster_recovery_plan.pdf",
        "docs/distributed_deployment.md": "docs/distributed_deployment.pdf",
        "docs/docker_connection.md": "docs/docker_connection.pdf",
        "docs/Env_Configuration.md": "docs/Env_Configuration.pdf",
        "docs/Final_Report.md": "docs/Final_Report.pdf",
        "docs/Global_Index.md": "docs/Global_Index.pdf",
        "docs/Header_Footer_Antigravity.md": "docs/Header_Footer_Antigravity.pdf",
        "docs/Header_Footer_Gemini.md": "docs/Header_Footer_Gemini.pdf",
        "docs/Installation_Guide.md": "docs/Installation_Guide.pdf",
        "docs/Operator_Installation_Guide.md": "docs/Operator_Installation_Guide.pdf",
        "docs/Presentation_Technical.md": "docs/Presentation_Technical.pdf",
        "docs/Presentation_User.md": "docs/Presentation_User.pdf",
        "docs/project_report.md": "docs/project_report.pdf",
        "docs/Recommendations.md": "docs/Recommendations.pdf",
        "docs/retro_planning.md": "docs/retro_planning.pdf",
        "docs/Scrum_Artifacts.md": "docs/Scrum_Artifacts.pdf",
        "docs/Start_Stop_Procedures.md": "docs/Start_Stop_Procedures.pdf",
        "docs/taiga_audit_report.md": "docs/taiga_audit_report.pdf",
        "docs/Technical_Document.md": "docs/Technical_Document.pdf",
        "docs/Uninstall_Guide.md": "docs/Uninstall_Guide.pdf",
        "docs/URL_Formats.md": "docs/URL_Formats.pdf",
        "docs/User_Description.md": "docs/User_Description.pdf",
        "docs/User_Guide.md": "docs/User_Guide.pdf",
        "docs/Wiki_Home.md": "docs/Wiki_Home.pdf",
        "docs/WSL_Deployment.md": "docs/WSL_Deployment.pdf",
        "docs/Virtualbox_Deployment.md": "docs/Virtualbox_Deployment.pdf",
        "docs/Hyper-V_Deployment.md": "docs/Hyper-V_Deployment.pdf",
        "docs/zabbix_monitoring.md": "docs/zabbix_monitoring.pdf",
        "docs/Logs_information.md": "docs/Logs_information.pdf",
        "docs/How_to_change_webhooks_and_emails.md": "docs/How_to_change_webhooks_and_emails.pdf",
        "docs/Submission_Requirements.md": "docs/Submission_Requirements.pdf",
        "README.md": "Readme.pdf"
    }

    print("Starting compilation of all system documents to PDF...")
    for md_rel, pdf_rel in target_mappings.items():
        md_abs = os.path.join(workspace, md_rel.replace("/", os.sep))
        pdf_abs = os.path.join(workspace, pdf_rel.replace("/", os.sep))
        
        if not os.path.exists(md_abs):
            fallback_abs = os.path.join(docs_dir, os.path.basename(md_abs))
            if os.path.exists(fallback_abs):
                md_abs = fallback_abs
            else:
                print(f"[WARNING] Source file not found: {md_rel} (skipping)")
                continue

        try:
            compiler = PDFCompiler(md_abs, pdf_abs)
            compiler.compile()
        except Exception as e:
            print(f"[ERROR] Failed to compile {md_rel}: {e}")

    print("[SUCCESS] PDF compilation sequence completed!")

    print("Starting compilation of Word DOCX documents...")
    docx_mappings = {md: pdf.replace('.pdf', '.docx') for md, pdf in target_mappings.items()}
    for md_rel, docx_rel in docx_mappings.items():
        md_abs = os.path.join(workspace, md_rel.replace("/", os.sep))
        docx_abs = os.path.join(workspace, docx_rel.replace("/", os.sep))
        
        if not os.path.exists(md_abs):
            fallback_abs = os.path.join(docs_dir, os.path.basename(md_abs))
            if os.path.exists(fallback_abs):
                md_abs = fallback_abs
            else:
                print(f"[WARNING] Source file not found for DOCX: {md_rel} (skipping)")
                continue

        try:
            compile_to_docx(md_abs, docx_abs)
        except Exception as e:
            print(f"[ERROR] Failed to compile DOCX {md_rel}: {e}")
    print("[SUCCESS] DOCX compilation sequence completed!")

if __name__ == "__main__":
    compile_all_documents()
